"""Converter modüllerinin birim testleri."""

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import pytest
from PIL import Image

from converters import ConversionError
from converters.excel_converter import convert_csv_to_xlsx, convert_xlsx_to_csv
from converters.html_converter import convert_html_to_pdf
from converters.image_converter import convert_image
from converters.pdf_converter import (
    convert_pdf_to_docx,
    convert_pdf_to_html,
    convert_pdf_to_images,
    convert_pdf_to_txt,
)
from converters.text_converter import convert_docx_to_txt, convert_txt_to_docx
from converters.word_to_pdf_converter import convert_docx_to_pdf


# --------------------------------------------------------------------------- #
# Görsel
# --------------------------------------------------------------------------- #
def _make_image(path, mode="RGBA", size=(8, 8), fmt="PNG"):
    Image.new(mode, size, (255, 0, 0, 255)).save(path, format=fmt)
    return path


def test_image_png_to_jpg(tmp_path):
    src = _make_image(str(tmp_path / "a.png"))
    dst = str(tmp_path / "a.jpg")
    convert_image(src, dst, "jpg")
    with Image.open(dst) as img:
        assert img.format == "JPEG"


def test_image_png_to_pdf(tmp_path):
    src = _make_image(str(tmp_path / "a.png"))
    dst = str(tmp_path / "a.pdf")
    convert_image(src, dst, "pdf")
    assert open(dst, "rb").read(4) == b"%PDF"


def test_image_gif_to_png(tmp_path):
    # GIF girdisi artık destekleniyor
    src = _make_image(str(tmp_path / "a.gif"), mode="RGB", fmt="GIF")
    dst = str(tmp_path / "a.png")
    convert_image(src, dst, "png")
    with Image.open(dst) as img:
        assert img.format == "PNG"


def test_image_tiff_to_jpg(tmp_path):
    src = _make_image(str(tmp_path / "a.tiff"), mode="RGB", fmt="TIFF")
    dst = str(tmp_path / "a.jpg")
    convert_image(src, dst, "jpg")
    with Image.open(dst) as img:
        assert img.format == "JPEG"


def test_image_rgba_to_jpg_loses_alpha(tmp_path):
    # Şeffaf görsel → JPG: RGB'ye çevrilmeli, hata vermemeli
    src = _make_image(str(tmp_path / "a.png"), mode="RGBA")
    dst = str(tmp_path / "a.jpg")
    convert_image(src, dst, "jpg")
    with Image.open(dst) as img:
        assert img.format == "JPEG"


def test_image_invalid_file_raises(tmp_path):
    src = tmp_path / "fake.png"
    src.write_bytes(b"not an image at all")
    with pytest.raises(ConversionError):
        convert_image(str(src), str(tmp_path / "out.jpg"), "jpg")


# --------------------------------------------------------------------------- #
# Metin / Word
# --------------------------------------------------------------------------- #
def test_txt_to_docx_roundtrip(tmp_path):
    src = tmp_path / "a.txt"
    src.write_text("merhaba dunya\nikinci satir\n", encoding="utf-8")
    docx_path = str(tmp_path / "a.docx")
    convert_txt_to_docx(str(src), docx_path)

    out = str(tmp_path / "out.txt")
    convert_docx_to_txt(docx_path, out)
    lines = open(out, encoding="utf-8").read().strip().splitlines()
    assert "merhaba dunya" in lines
    assert "ikinci satir" in lines


def test_docx_to_txt_includes_tables(tmp_path):
    import docx

    doc = docx.Document()
    doc.add_paragraph("baslik")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "hucre1"
    table.cell(0, 1).text = "hucre2"
    src = str(tmp_path / "a.docx")
    doc.save(src)

    out = str(tmp_path / "out.txt")
    convert_docx_to_txt(src, out)
    content = open(out, encoding="utf-8").read()
    assert "baslik" in content
    assert "hucre1" in content and "hucre2" in content


def test_txt_utf8_bom(tmp_path):
    # BOM'lu dosya da sorunsuz dönüşmeli
    src = tmp_path / "bom.txt"
    src.write_bytes(b"\xef\xbb\xbfmerhaba\n")
    docx_path = str(tmp_path / "bom.docx")
    convert_txt_to_docx(str(src), docx_path)
    assert os.path.getsize(docx_path) > 0


# --------------------------------------------------------------------------- #
# Excel
# --------------------------------------------------------------------------- #
def test_xlsx_to_csv(tmp_path):
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["ad", "soyad"])
    ws.append(["Ali", "Yilmaz"])
    src = str(tmp_path / "a.xlsx")
    wb.save(src)

    out = str(tmp_path / "a.csv")
    convert_xlsx_to_csv(src, out)
    content = open(out, encoding="utf-8").read()
    assert "Ali" in content and "Yilmaz" in content


def test_csv_to_xlsx(tmp_path):
    import csv
    import openpyxl

    src = tmp_path / "a.csv"
    with open(src, "w", newline="", encoding="utf-8") as f:
        csv.writer(f).writerows([["ad"], ["Ayse"]])

    out = str(tmp_path / "a.xlsx")
    convert_csv_to_xlsx(str(src), out)
    wb = openpyxl.load_workbook(out)
    assert wb.active["A2"].value == "Ayse"


# --------------------------------------------------------------------------- #
# HTML → PDF
# --------------------------------------------------------------------------- #
def test_html_to_pdf(tmp_path):
    src = tmp_path / "sayfa.html"
    src.write_text(
        "<html><body><h1>Merhaba</h1><p>Deneme metni</p></body></html>",
        encoding="utf-8",
    )
    out = str(tmp_path / "sayfa.pdf")
    convert_html_to_pdf(str(src), out)
    data = open(out, "rb").read()
    assert data[:4] == b"%PDF"
    assert len(data) > 100


# --------------------------------------------------------------------------- #
# PDF
# --------------------------------------------------------------------------- #
def _make_pdf(path, pages=1):
    import fitz

    doc = fitz.open()
    for _ in range(pages):
        page = doc.new_page()
        page.insert_text((72, 72), "Convrt test sayfasi")
    doc.save(path)
    doc.close()
    return str(path)


def test_pdf_to_txt(tmp_path):
    src = _make_pdf(tmp_path / "a.pdf")
    out = str(tmp_path / "a.txt")
    convert_pdf_to_txt(src, out)
    assert "Convrt test sayfasi" in open(out, encoding="utf-8").read()


def test_pdf_to_html(tmp_path):
    src = _make_pdf(tmp_path / "a.pdf")
    out = str(tmp_path / "a.html")
    convert_pdf_to_html(src, out)
    content = open(out, encoding="utf-8").read()
    assert "<html" in content and "Convrt test sayfasi" in content


def test_pdf_to_images_single_page(tmp_path):
    src = _make_pdf(tmp_path / "a.pdf", pages=1)
    images = convert_pdf_to_images(src, str(tmp_path), "png")
    assert len(images) == 1


def test_pdf_to_images_multiple_pages(tmp_path):
    src = _make_pdf(tmp_path / "a.pdf", pages=3)
    images = convert_pdf_to_images(src, str(tmp_path), "png")
    assert len(images) == 3


def test_pdf_to_docx(tmp_path):
    src = _make_pdf(tmp_path / "a.pdf")
    out = str(tmp_path / "a.docx")
    convert_pdf_to_docx(src, out)
    assert open(out, "rb").read(2) == b"PK"


def test_docx_to_pdf(tmp_path):
    import docx

    doc = docx.Document()
    doc.add_paragraph("Word'den PDF'e deneme")
    src = str(tmp_path / "a.docx")
    doc.save(src)

    out = str(tmp_path / "a.pdf")
    convert_docx_to_pdf(src, out)
    data = open(out, "rb").read()
    assert data[:4] == b"%PDF"
    assert len(data) > 100
