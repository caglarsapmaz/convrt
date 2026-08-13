import docx


def convert_docx_to_txt(input_path: str, output_path: str) -> None:
    """Word belgesindeki paragrafları ve tabloları düz metne dönüştürür."""
    doc = docx.Document(input_path)
    lines = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append("\t".join(cells))
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines) + "\n")


def convert_txt_to_docx(input_path: str, output_path: str) -> None:
    """Düz metni satır satır Word belgesine dönüştürür.

    utf-8-sig ile okunur; böylece BOM (byte order mark) içeren dosyalar da
    hatasız işlenir.
    """
    doc = docx.Document()
    with open(input_path, "r", encoding="utf-8-sig") as f:
        for line in f:
            doc.add_paragraph(line.strip())
    doc.save(output_path)
